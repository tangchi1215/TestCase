# %%
import os
import glob
import pandas as pd
import datetime
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


# %%
# 讀取Excel文件並進行初步處理
def load_and_prepare_data(file_path):
    try:
        df = pd.read_excel(file_path, engine='openpyxl')
        selected_columns = df[['功能類別', '測試個案編號', '個案說明', '預期結果', '測試日期', '測試結果', '備註']]
    except KeyError as e:
        print("標頭不符合指定格式")

    renamed_columns = selected_columns.rename(columns={
        '功能類別': '功能\n類別',
        '測試個案編號': '測試個案\n編號',
        '測試結果': '測試\n結果'
    })
    cleaned_data = renamed_columns.dropna(how='all')

    current_date = datetime.datetime.now()
    formatted_date = current_date.strftime('%Y/%m/%d')

    cleaned_data['測試日期'] = formatted_date
    cleaned_data['測試\n結果'] = '通過'
    base_name = file_path.split('.')[0]
    cleaned_data['測試個案\n編號'] = [f"{base_name}-{i + 1:02}" for i in range(len(cleaned_data))]
    return cleaned_data.fillna('')


# 設定單元格字體
def set_cell_font(cell, font_name='標楷體', font_size=Pt(12)):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = font_size
        if not paragraph.runs:
            run = paragraph.add_run()
            run.font.name = font_name
            run.font.size = font_size


def set_cell_width(cell, width):
    properties = cell._tc.get_or_add_tcPr()
    cell_width = OxmlElement('w:tcW')
    cell_width.set(qn('w:w'), str(width))
    cell_width.set(qn('w:type'), 'dxa')
    properties.append(cell_width)


def set_cell_center(cell):
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_vertical_alignment(cell, align=WD_ALIGN_VERTICAL.CENTER):
    cell.vertical_alignment = align


# 創建並格式化表格
def create_and_format_table(doc, data, column_widths):
    for index, row in data.iterrows():
        table = doc.add_table(rows=1, cols=len(data.columns))
        # 設定整行的列寬
        for idx, cell in enumerate(table.rows[0].cells):
            set_cell_width(cell, column_widths[idx].emu)

        set_table_header(table.rows[0], data.columns)
        add_table_data(table, row)
        doc.add_paragraph()  # 添加空行作為分隔


def set_table_header(row, columns):
    for i, col_name in enumerate(columns):
        row.cells[i].text = str(col_name)
        set_cell_shading(row.cells[i], 'D3D3D3')  # 灰色背景
        set_cell_border(row.cells[i])
        set_cell_font(row.cells[i])  # 設置字體
        set_cell_center(row.cells[i])
        set_vertical_alignment(row.cells[i])


def add_table_data(table, row_data):
    row_cells = table.add_row().cells
    for i, value in enumerate(row_data):
        row_cells[i].text = str(value)
        set_cell_border(row_cells[i])
        set_cell_font(row_cells[i])  # 設置字體
        set_cell_center(row_cells[i])
        set_vertical_alignment(row_cells[i])


def set_cell_shading(cell, fill_color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_border(cell, color='000000', size='4'):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


# %%
def scan_xlsx_files(directory):
    path_pattern = os.path.join(directory, '*.xlsx')
    files = glob.glob(path_pattern)
    return files


def display_files(files):
    if not files:
        print("找不到任何 .xlsx 檔案QAQ")
        return None

    print("以下為找到的 .xlsx 檔案：")
    for idx, file in enumerate(files, start=1):
        print(f"{idx}. {os.path.basename(file)}")

    return files


# %%
def user_select_file(files):
    try:
        selection = int(input("請輸入文件編號（ex:1, 2 ...）: "))
        if 1 <= selection <= len(files):
            return files[selection - 1]
        else:
            print("超出範圍，重來")
    except ValueError:
        print("無效的數值 = =")

    return None


# %%
# 主流程
def main():
    directory = input("請輸入測試案例.xlsx放的路徑: ")
    files = scan_xlsx_files(directory)
    displayed_files = display_files(files)
    if displayed_files:
        selected_file = user_select_file(displayed_files)
        if selected_file:
            print(f"你要產的測試報告：{os.path.basename(selected_file)}")
            path = selected_file.split('\\')[-1]
            cleaned_data = load_and_prepare_data(path)
            doc = Document('../docs/project_doc.docx')
            style = doc.styles['Normal']
            style.font.name = '標楷體'
            style.font.size = Pt(12)
            style._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
            doc._body.clear_content()

            column_widths = (
                Inches(0.8), Inches(0.44), Inches(1.962), Inches(1.24), Inches(0.2), Inches(0.3), Inches(0.2))
            create_and_format_table(doc, cleaned_data, column_widths)

            output_path = (path).split('.')[0] + '.docx'
            doc.save(output_path)
            print("產製完成~~!")
        else:
            print("你沒選 = =")


# %%
# 執行主程式
if __name__ == '__main__':
    main()
