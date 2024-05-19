from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


class TableFormatter:
    @staticmethod
    def set_cell_font(cell, font_name='標楷體', font_size=Pt(10)):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = font_size
            if not paragraph.runs:
                run = paragraph.add_run()
                run.font.name = font_name
                run.font.size = font_size

    @staticmethod
    def set_cell_width(cell, width):
        properties = cell._tc.get_or_add_tcPr()
        cell_width = OxmlElement('w:tcW')
        cell_width.set(qn('w:w'), str(width))
        cell_width.set(qn('w:type'), 'dxa')
        properties.append(cell_width)

    @staticmethod
    def set_cell_center(cell):
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    @staticmethod
    def set_vertical_alignment(cell, align=WD_ALIGN_VERTICAL.CENTER):
        cell.vertical_alignment = align

    @staticmethod
    def set_cell_shading(cell, fill_color):
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill_color)
        cell._tc.get_or_add_tcPr().append(shading_elm)

    @staticmethod
    def set_cell_border(cell, color='000000', size='4'):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), size)
            border.set(qn('w:color'), color)
            tcBorders.append(border)
        tcPr.append(tcBorders)

    @staticmethod
    def set_table_header(row, columns):
        for i, col_name in enumerate(columns):
            row.cells[i].text = str(col_name)
            TableFormatter.set_cell_shading(row.cells[i], 'D3D3D3')  # 灰色背景
            TableFormatter.set_cell_border(row.cells[i])
            TableFormatter.set_cell_font(row.cells[i])  # 設置字體
            TableFormatter.set_cell_center(row.cells[i])
            TableFormatter.set_vertical_alignment(row.cells[i])

    @staticmethod
    def create_and_format_table(doc, data, column_widths):
        for index, row in data.iterrows():
            table = doc.add_table(rows=1, cols=len(data.columns))
            # 設定整行的列寬
            for idx, cell in enumerate(table.rows[0].cells):
                TableFormatter.set_cell_width(cell, column_widths[idx].emu)

            TableFormatter.set_table_header(table.rows[0], data.columns)
            TableFormatter.add_table_data(table, row)
            doc.add_paragraph()  # 添加空行作為分隔

    @staticmethod
    def add_table_data(table, row_data):
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = str(value)
            TableFormatter.set_cell_border(row_cells[i])
            TableFormatter.set_cell_font(row_cells[i])  # 設置字體
            TableFormatter.set_cell_center(row_cells[i])
            TableFormatter.set_vertical_alignment(row_cells[i])
